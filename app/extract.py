"""
File extraction: turn an uploaded requirements file (.docx, .xlsx/.xlsm, .pdf, .txt/.csv)
into plain text for the Claude prompt.
"""
from pathlib import Path
import io


def extract_text(filename: str, raw_bytes: bytes) -> str:
    ext = Path(filename).suffix.lower().lstrip(".")

    if ext in ("txt", "csv"):
        return raw_bytes.decode("utf-8", errors="ignore")

    if ext == "docx":
        return _extract_docx(raw_bytes)

    if ext in ("xlsx", "xlsm"):
        return _extract_xlsx(raw_bytes)

    if ext == "pdf":
        return _extract_pdf(raw_bytes)

    # Fallback: try as text
    return raw_bytes.decode("utf-8", errors="ignore")


def _extract_docx(raw_bytes: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(raw_bytes))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_xlsx(raw_bytes: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(raw_bytes), data_only=True, read_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"--- Sheet: {ws.title} ---")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_pdf(raw_bytes: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw_bytes))
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n".join(parts)
